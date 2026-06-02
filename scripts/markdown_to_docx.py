#!/usr/bin/env python3
"""Convert casusbijeenkomst/LME-style markdown to a Word document."""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Installeer python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def add_runs_with_bold(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and "|" in lines[i]:
        line = lines[i].strip()
        if re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if cells:
            rows.append(cells)
        i += 1
    return rows, i


def heading_level(line: str) -> int | None:
    m = re.match(r"^(#{1,3})\s+(.+)$", line.strip())
    if m:
        return len(m.group(1))
    t = line.strip()
    if t in ("Leerdoelen", "Samenvatting"):
        return 1
    if re.match(r"^Tabel\s+\d+\.", t, re.I):
        return 2
    return None


def is_short_heading(line: str) -> bool:
    t = line.strip()
    if not t or len(t) > 90:
        return False
    if t.endswith((".", "!", "?")):
        return False
    if t.startswith(("-", "*", "|", "#")):
        return False
    return bool(re.match(r"^[A-ZÀ-ÖØ-Þ]", t))


def markdown_to_docx(markdown: str, out_path: Path, title: str | None = None) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown.replace("\r\n", "\n").split("\n")
    i = 0
    first_heading_done = False

    if title:
        doc.add_heading(title, level=1)
        first_heading_done = True

    bullet_buf: list[str] = []

    def flush_bullets() -> None:
        nonlocal bullet_buf
        for item in bullet_buf:
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, item)
        bullet_buf = []

    while i < len(lines):
        line = lines[i]
        trimmed = line.strip()

        if not trimmed:
            i += 1
            continue

        if trimmed.startswith("|") and "|" in trimmed[1:]:
            flush_bullets()
            rows, i = parse_table(lines, i)
            if len(rows) >= 1:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(cols):
                        cell_text = row[ci] if ci < len(row) else ""
                        table.rows[ri].cells[ci].text = cell_text
            continue

        if re.match(r"^[-*•]\s+", trimmed) or re.match(r"^\d+\.\s+", trimmed):
            text = re.sub(r"^[-*•]\s+", "", trimmed)
            text = re.sub(r"^\d+\.\s+", "", text)
            bullet_buf.append(text)
            i += 1
            continue

        flush_bullets()

        level = heading_level(trimmed)
        if level is None and is_short_heading(trimmed):
            level = 2 if first_heading_done else 1

        if level is not None:
            text = re.sub(r"^#{1,3}\s+", "", trimmed)
            doc.add_heading(text, level=min(level, 3))
            first_heading_done = True
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs_with_bold(p, trimmed)
        i += 1

    flush_bullets()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    if len(sys.argv) < 3:
        print("Usage: python markdown_to_docx.py <input.md> <output.docx> [title]", file=sys.stderr)
        sys.exit(1)
    md_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    title = sys.argv[3] if len(sys.argv) > 3 else None
    text = md_path.read_text(encoding="utf-8")
    markdown_to_docx(text, out_path, title=title)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
