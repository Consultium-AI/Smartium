#!/usr/bin/env python3
"""
Exporteert alle LME-samenvattingen naar Word (.docx), gegroepeerd per week binnen een blok.

Bron van waarheid voor welke LME bij welke week hoort: dezelfde structuur als
src/pages/SummaryPage.jsx (courseStructure).

Tekst wordt uit de JSX-bron gehaald in dezelfde leesvolgorde als de site/PDF: eerst
SummaryLayout-kop (title, description, caseLabel), daarna alle secties met h2–h4,
kaarten, lijsten en tabellen. Geen React-runtime nodig.

Gebruik:
  pip install python-docx
  python scripts/export_summaries_to_docx.py

Output (standaard): output/summary_exports/Blok3_Week_1.docx, Blok4_Week_1.docx, Blok5_Week_1.docx, ...

Blok 5 (en andere blokken met volledige lme-id's) worden automatisch uit courseStructure
in SummaryPage.jsx gelezen; blok 3/4 staan nog handmatig in COURSE_WEEKS.
"""

from __future__ import annotations

import re
import sys
from collections import defaultdict
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Installeer python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parents[1]
PAGES_DIR = ROOT / "src" / "pages"
SUMMARY_PAGE = PAGES_DIR / "SummaryPage.jsx"
SUMMARY_SUBDIR = PAGES_DIR / "summary"
SUMMARIES_DIR = ROOT / "src" / "summaries"


def as_extended_path(path: Path) -> Path:
    """Windows MAX_PATH (260): langere paden via \\\\?\\ prefix."""
    if sys.platform != "win32":
        return path
    resolved = path.resolve()
    text = str(resolved)
    if len(text) >= 260 and not text.startswith("\\\\?\\"):
        return Path("\\\\?\\" + text)
    return resolved


def path_is_file(path: Path) -> bool:
    return as_extended_path(path).is_file()


def read_text_file(path: Path) -> str:
    return as_extended_path(path).read_text(encoding="utf-8")


# Zelfde hiërarchie als courseStructure in SummaryPage.jsx (handmatig gesynchroniseerd).
COURSE_WEEKS: list[tuple[str, str, str, list[tuple[str, str, str]]]] = [
    (
        "blok3",
        "Blok 3: Groei en Ontwikkeling II",
        "Week 1",
        [
            (
                "Casus 1: Van embryo naar baby",
                "embryogenese",
                "Embryogenese Bouwplan",
            ),
            (
                "Casus 1: Van embryo naar baby",
                "beeldvorming",
                "Foetale beeldvorming en de 13 wekenecho",
            ),
            (
                "Casus 1: Van embryo naar baby",
                "extraembryonaal",
                "Het embryo en de extra embryonale structuren",
            ),
            (
                "Casus 2: Spina bifida",
                "extremiteiten",
                "Ontwikkeling extremiteiten en gewrichtsleer",
            ),
        ],
    ),
    (
        "blok4",
        "Blok 4: Microbiologie en Immunologie",
        "Week 1",
        [
            (
                "Casus 1: Algemene introductie",
                "aangeboren-immuniteit-herkenning",
                "Aangeboren immuniteit - herkenning",
            ),
            (
                "Casus 1: Algemene introductie",
                "introductie-bacteriologie",
                "Introductie bacteriologie",
            ),
            (
                "Casus 1: Algemene introductie",
                "introductie-immunologie",
                "Introductie en overzicht immunologie",
            ),
            (
                "Casus 1: Algemene introductie",
                "introductie-mycologie",
                "Introductie mycologie",
            ),
            (
                "Casus 1: Algemene introductie",
                "introductie-parasitologie",
                "Introductie parasitologie",
            ),
            (
                "Casus 1: Algemene introductie",
                "introductie-virologie",
                "Introductie virologie – deel 1",
            ),
            (
                "Casus 1: Algemene introductie",
                "voorbereiding-vow-microbiologie",
                "Voorbereiding VOW Microbiologie",
            ),
            (
                "Casus 2: Vrouw met Urineweginfectie",
                "aangeboren-immuniteit-respons",
                "Aangeboren immuniteit – respons",
            ),
            (
                "Casus 2: Vrouw met Urineweginfectie",
                "antibiotica-leerlijn",
                "Antibiotica leerlijn",
            ),
            (
                "Casus 2: Vrouw met Urineweginfectie",
                "introductie-antimicrobiele-therapie",
                "Introductie antimicrobiële therapie",
            ),
            (
                "Casus 2: Vrouw met Urineweginfectie",
                "urineweginfecties",
                "Urineweginfecties",
            ),
        ],
    ),
    (
        "blok4",
        "Blok 4: Microbiologie en Immunologie",
        "Week 2",
        [
            (
                "Casus 3: Kind met RS-virus",
                "infectiepreventie",
                "Infectiepreventie",
            ),
            (
                "Casus 3: Kind met RS-virus",
                "introductie-luchtweginfecties",
                "Introductie luchtweginfecties",
            ),
            (
                "Casus 3: Kind met RS-virus",
                "microscopische-anatomie-luchtwegen",
                "Microscopische anatomie luchtwegen",
            ),
            (
                "Casus 3: Kind met RS-virus",
                "pathofysiologie-virale-luchtweginfecties",
                "Pathofysiologie virale luchtweginfecties",
            ),
            (
                "Casus 3: Kind met RS-virus",
                "vervolg-introductie-virologie-2",
                "Vervolg introductie virologie 2",
            ),
            (
                "Casus 3: Kind met RS-virus",
                "virale-diagnostiek",
                "Virale diagnostiek",
            ),
            (
                "Casus 4: Kind met koorts",
                "cytokinen",
                "Cytokinen",
            ),
            (
                "Casus 4: Kind met koorts",
                "infectieuze-oorzaken-koorts",
                "Infectieuze oorzaken van koorts bij kinderen",
            ),
            (
                "Casus 4: Kind met koorts",
                "niet-infectieuze-oorzaken-koorts",
                "Niet-infectieuze oorzaken van koorts bij kinderen",
            ),
            (
                "Casus 4: Kind met koorts",
                "ontstekingsmediatoren",
                "Ontstekingsmediatoren",
            ),
            (
                "Casus 4: Kind met koorts",
                "wat-is-koorts",
                "Wat is koorts",
            ),
        ],
    ),
    (
        "blok4",
        "Blok 4: Microbiologie en Immunologie",
        "Week 3",
        [
            (
                "Casus 5: Onderste luchtweginfectie",
                "specifieke-verwekkers-lagere-luchtweginfecties",
                "Specifieke verwekkers lagere luchtweginfecties",
            ),
            (
                "Casus 5: Onderste luchtweginfectie",
                "therapie-lagere-luchtweginfecties",
                "Therapie lagere luchtweginfecties",
            ),
            (
                "Casus 6: Lymfeklier",
                "homing-migratie-recirculatie",
                "Homing, migratie en recirculatie",
            ),
            (
                "Casus 6: Lymfeklier",
                "infectieuze-lymfadenopathie",
                "Infectieuze lymfadenopathie – tweede lijn",
            ),
            (
                "Casus 6: Lymfeklier",
                "introductie-verworven-immuniteit",
                "Introductie verworven immuniteit",
            ),
            (
                "Casus 6: Lymfeklier",
                "secundaire-lymfoide-organen",
                "Secundaire lymfoïde organen",
            ),
        ],
    ),
    (
        "blok4",
        "Blok 4: Microbiologie en Immunologie",
        "Week 4",
        [
            ("Casus 8: HIV", "antivirale-therapie", "Antivirale therapie"),
            ("Casus 8: HIV", "de-grote-drie-malaria", "De grote drie Malaria"),
            (
                "Casus 8: HIV",
                "duurzame-hiv-zorg",
                "Duurzame hiv zorg in sub-Sahara-Afrika",
            ),
            (
                "Casus 8: HIV",
                "immunologische-consequenties-hiv",
                "Immunologische consequenties hiv",
            ),
            (
                "Casus 8: HIV",
                "tuberculose-bij-migranten",
                "Tuberculose bij migranten",
            ),
            (
                "Casus 8: HIV",
                "tuberculose-inleiding",
                "Tuberculose – een inleiding",
            ),
        ],
    ),
    (
        "blok4",
        "Blok 4: Microbiologie en Immunologie",
        "Week 5",
        [
            (
                "Casus 7",
                "lme5-schimmelinfecties",
                "Schimmelinfecties van de huid, nagels en haren",
            ),
            (
                "Casus 7",
                "lme6-voorbereiding-vow-milt",
                "Voorbereiding VOW Milt",
            ),
            (
                "Casus 9",
                "lme1-parasitaire-verwekkers-gastro-enteritis",
                "Parasitaire verwekkers van gastro-enteritis",
            ),
            (
                "Casus 9",
                "lme2-virale-verwekkers-gastro-enteritis",
                "Virale verwekkers van gastro-enteritis",
            ),
            (
                "Casus 9",
                "lme3-welk-antibioticum-kies-ik",
                "Welk antibioticum kies ik",
            ),
            (
                "Casus 10: Vaccinatie",
                "casus10-lme1-dwang-en-drang-historisch-perspectief",
                "Dwang en drang in historisch perspectief",
            ),
            (
                "Casus 10: Vaccinatie",
                "casus10-lme2-immunomodulatie",
                "Immunomodulatie",
            ),
            (
                "Casus 10: Vaccinatie",
                "casus10-lme3-rechtvaardiging-dwang-en-drang-morele-dilemmas",
                "Rechtvaardiging dwang en drang - Morele dilemma's bij vaccineren",
            ),
            (
                "Casus 10: Vaccinatie",
                "casus10-lme4-waarom-hoge-vaccinatiegraad-wiskunde-vaccinatie",
                "Waarom een hoge vaccinatiegraad? De wiskunde van vaccinatie",
            ),
        ],
    ),
    (
        "blok4",
        "Blok 4: Microbiologie en Immunologie",
        "Week 6",
        [
            (
                "Casus 11: Primaire immuundeficiënties",
                "casus11-lme1-leefstijl-en-immuunsysteem",
                "Leefstijl en immuunsysteem",
            ),
            (
                "Casus 12: Lijninfecties & resistenties",
                "casus12-lme1-antibioticaresistentie-en-therapie",
                "Antibioticaresistentie en therapie",
            ),
            (
                "Casus 12: Lijninfecties & resistenties",
                "casus12-lme2-sepsis",
                "Sepsis",
            ),
            (
                "Casus 12: Lijninfecties & resistenties",
                "casus12-lme3-patient-en-medicatieveiligheid",
                "Patient- en medicatieveiligheid",
            ),
            (
                "Casus 12: Lijninfecties & resistenties",
                "casus12-lme4-zorggerelateerde-infecties",
                "Zorggerelateerde infecties",
            ),
        ],
    ),
    (
        "blok4",
        "Blok 4: Microbiologie en Immunologie",
        "Week 7",
        [
            (
                "Casus 13: Een piloot met koorts",
                "casus13-lme1-antibiotica-introductie",
                "Antibiotica Introductie",
            ),
            (
                "Casus 13: Een piloot met koorts",
                "casus13-lme2-antibiotica-resistentie",
                "Antibiotica Resistentie",
            ),
        ],
    ),
]

# Blokken waarvan week/LME-lijst uit SummaryPage.jsx courseStructure komt.
AUTO_COURSE_BLOKS = ("blok5", "blok9", "blok10")


def parse_course_structure_blok(text: str, blok_id: str) -> list[tuple[str, str, str, list[tuple[str, str, str]]]]:
    """
    Haalt (blok_id, blok_name, week_name, [(case_name, lme_id, lme_title), ...]) uit courseStructure.
    """
    start = text.find(f"{blok_id}: {{")
    if start == -1:
        return []

    # Volgende top-level blok-sleutel na dit blok (blok9, blok10, …).
    next_blok = re.search(
        rf"\n\s{{4}}(blok\d+|embryo|week\d+|\w+):\s*\{{",
        text[start + len(blok_id) + 3 :],
    )
    end = start + len(blok_id) + 3 + next_blok.start() if next_blok else len(text)
    section = text[start:end]

    blok_name_m = re.search(r"""name:\s*["']([^"']+)["']""", section)
    blok_name = blok_name_m.group(1).strip() if blok_name_m else blok_id

    weeks: list[tuple[str, str, str, list[tuple[str, str, str]]]] = []
    # Quote-agnostisch: blok5/9 gebruiken dubbele quotes, blok10 enkele quotes.
    week_parts = re.split(r"""\{\s*name:\s*["'](Week \d+)["']\s*,\s*cases:\s*\[""", section)
    for i in range(1, len(week_parts), 2):
        week_name = week_parts[i]
        week_body = week_parts[i + 1]

        case_parts = re.split(r"""\{\s*name:\s*["'](Casus[^"']+)["']\s*,\s*lmes:\s*\[""", week_body)
        entries: list[tuple[str, str, str]] = []
        for j in range(1, len(case_parts), 2):
            case_name = case_parts[j]
            case_body = case_parts[j + 1]
            for m in re.finditer(
                r"""id:\s*(["'])(?P<id>(?:\\.|(?!\1).)*)\1\s*,\s*name:\s*(["'])(?P<name>(?:\\.|(?!\3).)*)\3\s*,\s*available:""",
                case_body,
            ):
                lme_id = m.group("id")
                lme_title = m.group("name").replace('\\"', '"').replace("\\'", "'")
                entries.append((case_name, lme_id, lme_title))

        weeks.append((blok_id, blok_name, week_name, entries))

    return weeks


def all_course_weeks() -> list[tuple[str, str, str, list[tuple[str, str, str]]]]:
    text = SUMMARY_PAGE.read_text(encoding="utf-8")
    auto: list[tuple[str, str, str, list[tuple[str, str, str]]]] = []
    for blok_id in AUTO_COURSE_BLOKS:
        auto.extend(parse_course_structure_blok(text, blok_id))
    return list(COURSE_WEEKS) + auto


def parse_summary_page_routing() -> tuple[dict[str, str], dict[str, Path]]:
    """lme_id -> component_name, component_name -> source file (relative to src/pages)."""
    text = SUMMARY_PAGE.read_text(encoding="utf-8")

    # Per if-blok: eerste <.../> is vaak <Header /> — we willen de *Summary-component.
    _chrome = frozenset({"Header", "BackButton", "Footer"})

    def pick_summary_component(chunk: str) -> str | None:
        matches = re.findall(r"<(\w+)\s*/>", chunk)
        for name in reversed(matches):
            if name.endswith("Summary"):
                return name
        for name in reversed(matches):
            if name not in _chrome:
                return name
        return matches[-1] if matches else None

    lme_to_component: dict[str, str] = {}
    parts = re.split(r"\n  if \(activeLme === ", text)
    for chunk in parts[1:]:
        m_id = re.match(r"'([^']+)'\)", chunk)
        if not m_id:
            continue
        lme_id = m_id.group(1)
        comp = pick_summary_component(chunk)
        if comp:
            lme_to_component[lme_id] = comp

    component_to_file: dict[str, Path] = {}

    def register(names: list[str], rel: str) -> None:
        p = Path(rel)
        for n in names:
            component_to_file[n] = p

    # Named imports uit ./summary/ — vaak meerdere regels: import {\n  Foo,\n  Bar,\n} from '...'
    for m_named in re.finditer(
        r"import\s+\{([^}]+)\}\s+from\s+['\"](\.\/summary\/[^'\"]+)['\"]",
        text,
    ):
        names = [x.strip() for x in m_named.group(1).split(",") if x.strip()]
        rel = "summary/" + m_named.group(2).split("/")[-1]
        if not rel.endswith(".jsx"):
            rel += ".jsx"
        register(names, rel)

    for line in text.splitlines():
        line = line.strip()
        m_default = re.match(
            r"import\s+(\w+)\s+from\s+['\"](\.\.\/summaries\/[^'\"]+)['\"]",
            line,
        )
        if m_default:
            name = m_default.group(1)
            rel_path = m_default.group(2).replace("../", "")
            component_to_file[name] = Path("..") / Path(rel_path).with_suffix(".jsx")
            continue

    return lme_to_component, component_to_file


def resolve_js_path(rel_from_pages: Path) -> Path:
    resolved = (PAGES_DIR / rel_from_pages).resolve()
    if not path_is_file(resolved):
        alt = resolved.with_suffix(".tsx")
        if path_is_file(alt):
            return alt
    return resolved


def extract_arrow_body(content: str, component_name: str) -> str | None:
    """
    Functiebody na de pijl. Ondersteunt zowel block-body `=> { … }` (blok9) als
    impliciete return met haakjes `=> ( <JSX/> )` (blok10).
    """
    needle = f"const {component_name}"
    idx = content.find(needle)
    if idx == -1:
        return None
    arrow = content.find("=>", idx)
    if arrow == -1:
        return None

    j = arrow + 2
    while j < len(content) and content[j] in " \t\r\n":
        j += 1
    if j >= len(content):
        return None

    open_ch = content[j]
    close_ch = {"{": "}", "(": ")"}.get(open_ch)
    if close_ch is None:
        return None

    depth = 0
    i = j
    while i < len(content):
        c = content[i]
        if c == open_ch:
            depth += 1
        elif c == close_ch:
            depth -= 1
            if depth == 0:
                return content[j + 1 : i]
        i += 1
    return None


def strip_jsx_exprs(s: str) -> str:
    out: list[str] = []
    i = 0
    n = len(s)
    while i < n:
        if i + 1 < n and s[i : i + 2] == "{{":
            out.append("{")
            i += 2
            continue
        if s[i] == "{":
            depth = 1
            i += 1
            while i < n and depth:
                if i + 1 < n and s[i : i + 2] == "{{":
                    i += 2
                    continue
                if s[i] == "{":
                    depth += 1
                elif s[i] == "}":
                    depth -= 1
                i += 1
            out.append(" ")
            continue
        out.append(s[i])
        i += 1
    return "".join(out)


def extract_summary_layout_parts(body: str) -> tuple[str, str]:
    """
    Haalt attributenstring en children van de eerste <SummaryLayout> uit component-body.
    Als er geen SummaryLayout is, valt terug op hele body.
    """
    m = re.search(r"<SummaryLayout([^>]*)>([\s\S]*?)</SummaryLayout>", body)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return "", body


def summary_layout_header_lines(attrs: str) -> str:
    """title, description, caseLabel uit SummaryLayout-attributen (zoals PDF-kop)."""
    lines: list[str] = []
    for attr in ("title", "description", "caseLabel"):
        for m in re.finditer(rf'{attr}="([^"]*)"', attrs):
            t = m.group(1).strip()
            if t:
                lines.append(t)
    return "\n\n".join(lines)


def jsx_linearize_to_text(jsx_inner: str) -> str:
    """
    Zet JSX-body om naar doorlopende tekst in bronvolgorde (zoals gerenderde pagina/PDF).
    Gebruikt markdown-achtige koppen voor Word: ## = h2, ### = h3, #### = h4.
    """
    s = jsx_inner
    s = re.sub(r"/\*[\s\S]*?\*/", "", s)
    s = re.sub(r"\{/\*[\s\S]*?\*/\}", "", s)
    s = strip_jsx_exprs(s)

    # Zelfsluitende React-componenten (Lucide, custom) — geen tekst
    s = re.sub(
        r"<[A-Z][A-Za-z0-9.]*(?:\s[^>]*)?/>",
        "",
        s,
    )
    # Zelfsluitende html-achtige tags
    s = re.sub(r"<[a-zA-Z][a-zA-Z0-9]*(?:\s[^>]*)?/>", "", s)

    inline_tags = r"</?(?:strong|b|em|i|mark|small|span|a)(?:\s[^>]*)?>"
    s = re.sub(inline_tags, "", s, flags=re.I)

    def blk(pat: str, rep: str, flags: int = 0) -> None:
        nonlocal s
        s = re.sub(pat, rep, s, flags=flags)

    blk(r"<h1[^>]*>", "\n\n# ", re.I)
    blk(r"</h1>", "\n", re.I)
    blk(r"<h2[^>]*>", "\n\n## ", re.I)
    blk(r"</h2>", "\n", re.I)
    blk(r"<h3[^>]*>", "\n\n### ", re.I)
    blk(r"</h3>", "\n", re.I)
    blk(r"<h4[^>]*>", "\n\n#### ", re.I)
    blk(r"</h4>", "\n", re.I)
    blk(r"<h5[^>]*>", "\n\n##### ", re.I)
    blk(r"</h5>", "\n", re.I)
    blk(r"<h6[^>]*>", "\n\n###### ", re.I)
    blk(r"</h6>", "\n", re.I)

    blk(r"<p[^>]*>", "\n", re.I)
    blk(r"</p>", "\n", re.I)
    blk(r"<li[^>]*>", "\n• ", re.I)
    blk(r"</li>", "\n", re.I)
    blk(r"<br\s*/?>", "\n", re.I)
    blk(r"<hr[^>]*/?>", "\n\n---\n", re.I)

    blk(r"<th[^>]*>", "\n[", re.I)
    blk(r"</th>", "] ", re.I)
    blk(r"<td[^>]*>", "\n", re.I)
    blk(r"</td>", "\t", re.I)
    blk(r"<tr[^>]*>", "\n", re.I)
    blk(r"</tr>", "\n", re.I)

    blk(r"<section[^>]*>", "\n", re.I)
    blk(r"</section>", "\n", re.I)
    blk(r"<article[^>]*>", "\n", re.I)
    blk(r"</article>", "\n", re.I)
    blk(r"<header[^>]*>", "\n", re.I)
    blk(r"</header>", "\n", re.I)
    blk(r"<footer[^>]*>", "\n", re.I)
    blk(r"</footer>", "\n", re.I)
    blk(r"<main[^>]*>", "\n", re.I)
    blk(r"</main>", "\n", re.I)

    # Overige tags: verwijderen, inhoud blijft
    s = re.sub(r"</[a-zA-Z][a-zA-Z0-9]*>", "\n", s)
    s = re.sub(r"<[a-zA-Z][a-zA-Z0-9]*(?:\s[^>]*)?>", " ", s)

    s = s.replace("&nbsp;", " ")
    s = s.replace("&#39;", "'")
    s = s.replace("&quot;", '"')
    s = re.sub(r"\n•\s*•\s*", "\n• ", s)
    s = re.sub(r"[ \t]+\n", "\n", s)
    s = re.sub(r"\n[ \t]+", "\n", s)
    s = re.sub(r"[ \t]{2,}", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    s = merge_markdown_heading_continuations(s)
    return s.strip()


def merge_markdown_heading_continuations(text: str) -> str:
    """Voegt losse '#…#' regels samen met de volgende regel (multiline h2 in JSX)."""
    lines = text.splitlines()
    out: list[str] = []
    i = 0
    while i < len(lines):
        cur = lines[i].strip()
        if re.match(r"^#{1,6}\s*$", cur) and i + 1 < len(lines):
            nxt = lines[i + 1].strip()
            if nxt and not nxt.startswith("#"):
                out.append(f"{cur} {nxt}")
                i += 2
                continue
        out.append(lines[i])
        i += 1
    return "\n".join(out)


def jsx_body_to_full_text(component_body: str) -> str:
    """tableOfContents + return buiten beschouwing laten: alleen SummaryLayout."""
    attrs, children = extract_summary_layout_parts(component_body)
    header = summary_layout_header_lines(attrs)
    body = jsx_linearize_to_text(children)
    if header and body:
        return f"{header}\n\n{body}"
    if header:
        return header
    return jsx_linearize_to_text(component_body)


def extract_titles_from_toc(jsx: str) -> list[str]:
    """Haalt title: '...' uit tableOfContents-array (eerste grote array na const tableOfContents)."""
    idx = jsx.find("const tableOfContents")
    if idx == -1:
        return []
    sub = jsx[idx : idx + 8000]
    return re.findall(r"title:\s*'((?:\\'|[^'])*)'", sub)


def ordered_image_imports(main_jsx: str) -> list[str]:
    """Bestandsnamen zoals Image01IntroductieSummary uit ./Image...."""
    return re.findall(r'from ["\']\./(Image\d+\w*Summary)["\']', main_jsx)


def ordered_sect_components(summary_content: str, comp: str) -> list[str]:
    """Sect01…/Section01… componenten in leesvolgorde uit SummaryLayout-children."""
    main_body = extract_arrow_body(summary_content, comp) or summary_content
    _, children = extract_summary_layout_parts(main_body)
    # blok9 gebruikt <Sect01…/>, blok10 gebruikt <Section01…/> (in ./sections/).
    return re.findall(r"<(Sect(?:ion)?\d+\w*)\s*/>", children)


def local_component_import_paths(content: str) -> dict[str, str]:
    """component-naam -> relatief import-pad (bv. './sections/Section01Titel')."""
    res: dict[str, str] = {}
    for m in re.finditer(
        r"""import\s+(\w+)\s+from\s+["'](\.[^"']+)["']""", content
    ):
        res[m.group(1)] = m.group(2)
    return res


def resolve_sect_file(folder: Path, sect: str, import_paths: dict[str, str]) -> Path:
    """Bepaal het .jsx-bestand voor een sectiecomponent via het import-pad, met fallbacks."""
    rel = import_paths.get(sect)
    if rel:
        candidate = (folder / rel)
        if candidate.suffix.lower() != ".jsx":
            candidate = candidate.with_suffix(".jsx")
        if path_is_file(candidate):
            return candidate
    # Fallbacks: zelfde map of ./sections/
    for cand in (folder / f"{sect}.jsx", folder / "sections" / f"{sect}.jsx"):
        if path_is_file(cand):
            return cand
    return folder / f"{sect}.jsx"


def extract_default_export_body(content: str) -> str | None:
    """Body na return (…) in export default function …()."""
    m = re.search(
        r"export\s+default\s+function\s+\w+\s*\([^)]*\)\s*\{([\s\S]*)\}\s*$",
        content.strip(),
    )
    if not m:
        return None
    inner = m.group(1)
    ret = re.search(r"return\s*\(\s*([\s\S]*)\s*\)\s*;?\s*$", inner.strip())
    return ret.group(1) if ret else inner


def _md_inline_to_text(t: str) -> str:
    """Verwijder markdown-vetmarkeringen (**…**) uit Inline/PBody-tekst."""
    return t.replace("**", "")


def _extract_string_list(segment: str) -> list[str]:
    """Alle string-literals (backtick/'/\") in bronvolgorde uit één rij."""
    cells: list[str] = []
    for m in re.finditer(
        r"`((?:\\`|[^`])*)`|'((?:\\'|[^'])*)'|\"((?:\\\"|[^\"])*)\"", segment
    ):
        val = m.group(1)
        if val is None:
            val = m.group(2)
        if val is None:
            val = m.group(3)
        cells.append(_md_inline_to_text(val or ""))
    return cells


def _datatable_array_to_html(arr_text: str) -> str:
    """[[kop…],[rij…],…] (JS-array) -> eenvoudige <table> voor de bestaande tabel-linearisatie."""
    rows = re.findall(r"\[((?:[^\[\]])*)\]", arr_text)
    if not rows:
        return ""
    html = ["<table>"]
    for ri, row in enumerate(rows):
        cells = _extract_string_list(row)
        tag = "th" if ri == 0 else "td"
        html.append("<tr>" + "".join(f"<{tag}>{c}</{tag}>" for c in cells) + "</tr>")
    html.append("</table>")
    return "\n" + "".join(html) + "\n"


def resolve_const_arrays(content: str) -> dict[str, str]:
    """const NAME = [ … ] -> arraytekst (balanced brackets), voor <DataTable rows={NAME} />."""
    res: dict[str, str] = {}
    for m in re.finditer(r"const\s+([A-Za-z_]\w*)\s*=\s*\[", content):
        name = m.group(1)
        start = content.index("[", m.end() - 1)
        depth = 0
        i = start
        while i < len(content):
            c = content[i]
            if c == "[":
                depth += 1
            elif c == "]":
                depth -= 1
                if depth == 0:
                    res[name] = content[start : i + 1]
                    break
            i += 1
    return res


def expand_section_shared_components(body: str, full_content: str) -> str:
    """
    Blok 10-secties houden tekst in SectionShared-componenten met template-literals:
      <Inline>{`tekst`}</Inline>, <PBody text={`…`} />, <DataTable rows={[…] of NAME} />.
    Zet die om naar gewone HTML vóór strip_jsx_exprs, zodat de tekst behouden blijft.
    """
    consts = resolve_const_arrays(full_content)

    def rows_ident(m: re.Match) -> str:
        arr = consts.get(m.group(1))
        return "rows={" + arr + "}" if arr else m.group(0)

    body = re.sub(r"rows=\{([A-Za-z_]\w*)\}", rows_ident, body)
    body = re.sub(
        r"<DataTable\s+rows=\{(\[[\s\S]*?\])\}\s*/>",
        lambda m: _datatable_array_to_html(m.group(1)),
        body,
    )

    def pbody(m: re.Match) -> str:
        txt = m.group(1)
        return "".join(
            f"<p>{_md_inline_to_text(line.strip())}</p>"
            for line in txt.split("\n")
            if line.strip()
        )

    body = re.sub(r"<PBody\s+text=\{`([\s\S]*?)`\}\s*/>", pbody, body)
    body = re.sub(
        r"<Inline>\s*\{`([\s\S]*?)`\}\s*</Inline>",
        lambda m: _md_inline_to_text(m.group(1)),
        body,
    )
    return body


def jsx_file_to_linear_text(content: str, comp_name: str | None = None) -> str:
    body = extract_arrow_body(content, comp_name) if comp_name else None
    if body is None:
        body = extract_default_export_body(content)
    if body is None:
        body = content
    body = expand_section_shared_components(body, content)
    return jsx_body_to_full_text(body)


def load_summary_with_parts(content: str, comp: str, folder: Path) -> str:
    main_body = extract_arrow_body(content, comp) or content
    attrs_wrap, _ = extract_summary_layout_parts(main_body)
    pre = summary_layout_header_lines(attrs_wrap)
    if pre:
        pre = pre + "\n\n---\n\n"

    images = ordered_image_imports(content)
    if images:
        toc_titles = extract_titles_from_toc(content)
        chunks: list[str] = []
        for i, mod in enumerate(images):
            title = toc_titles[i] if i < len(toc_titles) else f"Onderdeel {i + 1}"
            ip = folder / f"{mod}.jsx"
            if not path_is_file(ip):
                chunks.append(f"## {title}\n\n[Bestand ontbreekt: {ip.name}]")
                continue
            inner = read_text_file(ip)
            txt = jsx_file_to_linear_text(inner, mod)
            chunks.append(f"## {title}\n\n{txt}")
        return pre + "\n\n".join(chunks)

    sects = ordered_sect_components(content, comp)
    if sects:
        toc_titles = extract_titles_from_toc(content)
        import_paths = local_component_import_paths(content)
        chunks = []
        for i, sect in enumerate(sects):
            title = toc_titles[i] if i < len(toc_titles) else sect
            ip = resolve_sect_file(folder, sect, import_paths)
            if not path_is_file(ip):
                chunks.append(f"## {title}\n\n[Bestand ontbreekt: {ip.name}]")
                continue
            inner = read_text_file(ip)
            txt = jsx_file_to_linear_text(inner)
            chunks.append(f"## {title}\n\n{txt}")
        return pre + "\n\n".join(chunks)

    return jsx_body_to_full_text(main_body)


# SummaryPage importeert casus12-lme1 als namespace (* as …); geen directe component→pad mapping.
LME_FALLBACK_FILES: dict[str, tuple[str, Path]] = {
    "casus12-lme1-antibioticaresistentie-en-therapie": (
        "Casus12Lme1AntibioticaresistentieSummary",
        SUMMARIES_DIR
        / "casus12-lme1-antibioticaresistentie-en-therapie"
        / "Casus12Lme1AntibioticaresistentieSummary.jsx",
    ),
}


def load_main_module_text(
    lme_id: str,
    lme_to_component: dict[str, str],
    component_to_file: dict[str, Path],
) -> str:
    comp = lme_to_component.get(lme_id)
    if not comp:
        return f"[Geen route gevonden in SummaryPage.jsx voor lme={lme_id}]"

    path: Path | None = None
    if lme_id in LME_FALLBACK_FILES:
        comp_fb, path_fb = LME_FALLBACK_FILES[lme_id]
        comp = comp_fb
        path = path_fb
    else:
        rel = component_to_file.get(comp)
        if not rel:
            return f"[Geen import-pad voor component {comp}]"
        path = resolve_js_path(rel)

    if not path_is_file(path):
        return f"[Bestand ontbreekt: {path}]"

    content = read_text_file(path)
    return load_summary_with_parts(content, comp, path.parent)


def append_formatted_plaintext_to_doc(doc: Document, text: str) -> None:
    """Verwerkt regels met # … ###### koppen (output van jsx_linearize_to_text)."""
    buf: list[str] = []

    def flush() -> None:
        if not buf:
            return
        combined = "\n".join(buf).strip()
        if combined:
            doc.add_paragraph(combined)
        buf.clear()

    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            flush()
            continue
        if line.startswith("###### "):
            flush()
            doc.add_heading(line[7:].strip(), level=8)
            continue
        if line.startswith("##### "):
            flush()
            doc.add_heading(line[6:].strip(), level=7)
            continue
        if line.startswith("#### "):
            flush()
            doc.add_heading(line[5:].strip(), level=6)
            continue
        if line.startswith("### "):
            flush()
            doc.add_heading(line[4:].strip(), level=5)
            continue
        if line.startswith("## "):
            flush()
            doc.add_heading(line[3:].strip(), level=4)
            continue
        if line.startswith("# "):
            flush()
            doc.add_heading(line[2:].strip(), level=3)
            continue
        if line == "---":
            flush()
            continue
        buf.append(raw.rstrip())
    flush()


def build_week_documents(
    out_dir: Path,
) -> None:
    lme_to_component, component_to_file = parse_summary_page_routing()
    out_dir.mkdir(parents=True, exist_ok=True)

    # (blok_id, blok_name, week_name) -> …
    by_week: dict[tuple[str, str, str], list[tuple[str, str, str]]] = defaultdict(
        list
    )
    missing: list[str] = []

    for blok_id, blok_name, week_name, entries in all_course_weeks():
        for case_name, lme_id, lme_title in entries:
            text = load_main_module_text(lme_id, lme_to_component, component_to_file)
            if text.startswith("["):
                missing.append(f"{blok_id} {week_name} {lme_id}: {text}")
            by_week[(blok_id, blok_name, week_name)].append((case_name, lme_title, text))

    for (blok_id, blok_name, week_name), items in by_week.items():
        safe_week = week_name.replace(" ", "_")
        fname = f"{blok_id.capitalize()}_{safe_week}.docx"
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        doc.add_heading(f"{blok_name} — {week_name}", 0)
        doc.add_paragraph(
            "Automatisch geëxporteerd uit de Smartium JSX-samenvattingen."
        )

        current_case: str | None = None
        for case_name, lme_title, text in items:
            if case_name != current_case:
                current_case = case_name
                doc.add_heading(case_name, level=1)

            doc.add_heading(f"LME: {lme_title}", level=2)
            append_formatted_plaintext_to_doc(doc, text)

        out_path = out_dir / fname
        doc.save(out_path)
        print(f"Geschreven: {out_path}")

    if missing:
        print("\nWaarschuwingen:", file=sys.stderr)
        for m in missing:
            print(f"  {m}", file=sys.stderr)


def main() -> None:
    out = ROOT / "output" / "summary_exports"
    if len(sys.argv) > 1:
        out = Path(sys.argv[1]).resolve()
    build_week_documents(out)


if __name__ == "__main__":
    main()
